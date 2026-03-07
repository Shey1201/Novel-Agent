"""
小说下载API
支持导出为Word文档格式
"""

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional, List
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import re
from urllib.parse import quote

router = APIRouter(prefix="/download", tags=["download"])


class ChapterData(BaseModel):
    """章节数据"""
    id: str
    title: str
    content: str


class NovelDownloadRequest(BaseModel):
    """小说下载请求"""
    novel_title: str
    chapters: List[ChapterData]
    download_type: str  # "full", "single", "range"
    start_chapter: Optional[int] = None  # 范围下载的起始章节（1-based）
    end_chapter: Optional[int] = None    # 范围下载的结束章节（1-based）
    single_chapter_id: Optional[str] = None  # 单章节下载的章节ID


def html_to_text(html_content: str) -> str:
    """将HTML内容转换为纯文本"""
    if not html_content:
        return ""
    
    # 先处理换行标签
    text = html_content.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
    # 处理段落标签
    text = text.replace('</p>', '\n\n').replace('<p>', '')
    # 移除其他HTML标签
    text = re.sub(r'<[^>]+>', '', text)
    # 转换HTML实体
    text = text.replace('&nbsp;', ' ')
    text = text.replace('&lt;', '<')
    text = text.replace('&gt;', '>')
    text = text.replace('&amp;', '&')
    text = text.replace('&quot;', '"')
    # 规范化空白字符
    text = re.sub(r'\n\s*\n', '\n\n', text)
    return text.strip()


def set_chinese_font(run, font_name='宋体', font_size=None):
    """设置中文字体"""
    run.font.name = font_name
    # 设置东亚字体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if font_size:
        run.font.size = Pt(font_size)


def create_word_document(novel_title: str, chapters: List[ChapterData]) -> BytesIO:
    """创建Word文档"""
    doc = Document()
    
    # 设置文档默认字体 - 使用中英文字体组合
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    
    # 添加小说标题
    title = doc.add_heading(level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run(novel_title)
    set_chinese_font(title_run, '黑体', 22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    
    # 添加空行
    doc.add_paragraph()
    
    # 添加章节内容
    for i, chapter in enumerate(chapters):
        # 章节标题
        chapter_heading = doc.add_heading(level=1)
        chapter_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        heading_run = chapter_heading.add_run(chapter.title)
        set_chinese_font(heading_run, '黑体', 16)
        heading_run.font.bold = True
        
        # 章节内容
        content_text = html_to_text(chapter.content)
        if content_text:
            paragraphs = content_text.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    p = doc.add_paragraph()
                    run = p.add_run(para_text.strip())
                    set_chinese_font(run, '宋体', 12)
                    p.paragraph_format.first_line_indent = Inches(0.5)
                    p.paragraph_format.line_spacing = 1.5
        else:
            # 如果内容为空，添加提示
            p = doc.add_paragraph()
            run = p.add_run("（本章暂无内容）")
            set_chinese_font(run, '宋体', 12)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.font.italic = True
        
        # 章节之间添加分页（除了最后一个章节）
        if i < len(chapters) - 1:
            doc.add_page_break()
    
    # 保存到内存
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


@router.post("/novel")
async def download_novel(request: NovelDownloadRequest):
    """
    下载小说为Word文档
    
    - download_type: "full"(全部), "single"(单章), "range"(范围)
    - 根据download_type选择相应的章节进行导出
    """
    try:
        # 根据下载类型筛选章节
        selected_chapters = []
        
        if request.download_type == "full":
            selected_chapters = request.chapters
            
        elif request.download_type == "single":
            if not request.single_chapter_id:
                raise HTTPException(status_code=400, detail="单章节下载需要提供章节ID")
            selected_chapters = [
                ch for ch in request.chapters 
                if ch.id == request.single_chapter_id
            ]
            if not selected_chapters:
                raise HTTPException(status_code=404, detail="未找到指定章节")
                
        elif request.download_type == "range":
            if request.start_chapter is None or request.end_chapter is None:
                raise HTTPException(status_code=400, detail="范围下载需要提供起始和结束章节")
            
            start_idx = max(0, request.start_chapter - 1)
            end_idx = min(len(request.chapters), request.end_chapter)
            
            if start_idx >= end_idx:
                raise HTTPException(status_code=400, detail="无效的章节范围")
                
            selected_chapters = request.chapters[start_idx:end_idx]
        else:
            raise HTTPException(status_code=400, detail="无效的下载类型")
        
        if not selected_chapters:
            raise HTTPException(status_code=400, detail="没有可下载的章节")
        
        # 生成Word文档
        doc_buffer = create_word_document(request.novel_title, selected_chapters)
        
        # 生成文件名
        safe_title = re.sub(r'[^\w\u4e00-\u9fa5]', '_', request.novel_title)
        timestamp = __import__('datetime').datetime.now().strftime('%Y%m%d')
        
        if request.download_type == "full":
            filename = f"{safe_title}_全文_{timestamp}.docx"
        elif request.download_type == "single":
            chapter_title = selected_chapters[0].title if selected_chapters else "章节"
            safe_chapter = re.sub(r'[^\w\u4e00-\u9fa5]', '_', chapter_title)
            filename = f"{safe_title}_{safe_chapter}_{timestamp}.docx"
        else:  # range
            filename = f"{safe_title}_第{request.start_chapter}-{request.end_chapter}章_{timestamp}.docx"
        
        # 返回文件流 - 使用 RFC 5987 编码处理中文文件名
        # 同时提供 filename (ASCII) 和 filename* (UTF-8) 以兼容不同浏览器
        ascii_filename = re.sub(r'[^\x00-\x7F]', '_', filename)  # 将非ASCII字符替换为下划线
        utf8_filename = quote(filename, safe='')
        
        return StreamingResponse(
            doc_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=\"{ascii_filename}\"; filename*=UTF-8''{utf8_filename}",
                "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        # 打印详细错误日志
        import traceback
        print(f"[Download Error] {str(e)}")
        print(f"[Download Error Traceback] {traceback.format_exc()}")
        
        # 将技术错误转换为用户友好的提示
        error_msg = str(e)
        if "codec" in error_msg.lower() or "encode" in error_msg.lower():
            user_friendly_msg = "文档编码错误，请检查小说内容是否包含特殊字符"
        elif "font" in error_msg.lower():
            user_friendly_msg = "字体设置错误，请重试或联系管理员"
        else:
            user_friendly_msg = f"生成文档失败: {str(e)[:50]}"
        raise HTTPException(status_code=500, detail=user_friendly_msg)
